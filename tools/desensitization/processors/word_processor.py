"""
Word 文档处理器（.docx）
基于 python-docx，遍历段落、表格、页眉页脚中的文本，
在 run 级别做替换以最大程度保留格式（字体、颜色、加粗等）。
"""
import logging
from typing import List

from docx import Document

from core.llm_client import Replacement
from processors.base import BaseProcessor

logger = logging.getLogger(__name__)


class WordProcessor(BaseProcessor):
    """Word .docx 处理器"""

    def extract_text(self, file_path: str) -> str:
        doc = Document(file_path)
        parts: List[str] = []
        for para in self._iter_all_paragraphs(doc):
            if para.text:
                parts.append(para.text)
        return "\n".join(parts)

    def apply_replacements(
        self, file_path: str, replacements: List[Replacement]
    ) -> List[Replacement]:
        doc = Document(file_path)
        applied_set = set()
        applied: List[Replacement] = []

        for para in self._iter_all_paragraphs(doc):
            self._replace_in_paragraph(para, replacements, applied_set, applied)

        doc.save(file_path)
        return applied

    @staticmethod
    def _iter_all_paragraphs(doc):
        """遍历文档中所有段落（正文 + 表格 + 页眉页脚）"""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                for para in header_footer.paragraphs:
                    yield para
